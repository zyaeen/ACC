import numpy as np
import matplotlib.pyplot as plt
import collections
import pandas as pd
from random import randrange
import os
import docx
from docx import document


n = 10
k = 24

argument = lambda i: k * i

# Задание отсчетов цен
c = [1.5 + 1 * np.sin(argument(i)) for i in range(1, n + 1)]
c2 = [2 + 1 * np.cos(argument(i)) for i in range(1, n + 1)]
# Случайная генерация объемов

B = [np.random.uniform(2, 5) for _ in range(n)]
B2 = [np.random.uniform(1, 7) for _ in range(n)]

Dict = {}
for i in range(n):
    Dict[c[i]] = [B[i], i + 1]

Dict2 = {}
for i in range(n):
    Dict2[c2[i]] = [B2[i], i + 1]

data = {}
for i in range(n):
    data['A{}'.format(i + 1)] = [
                                 '{:.3f}'.format(c[i]),
                                 '{:.3f}'.format(B[i])
                                ]

data2 = {}
for i in range(n):
    data2['B{}'.format(i + 1)] = [
                                 '{:.3f}'.format(c2[i]),
                                 '{:.3f}'.format(B2[i])
                                ]


doc = docx.Document()

parag = doc.add_paragraph('Генерация данных')


menuTable = doc.add_table(rows=1,cols=3)
menuTable.style= 'TableGrid'
hdr_Cells = menuTable.rows[0].cells
hdr_Cells[0].text = 'Указатель'
hdr_Cells[1].text = 'Значения p'
hdr_Cells[2].text = 'Значения α'


for A in data.keys():
    row_Cells = menuTable.add_row().cells
    row_Cells[0].text= str(A)
    row_Cells[1].text = str(data[A][0])
    row_Cells[2].text = str(data[A][1])

parag = doc.add_paragraph('     ')

menuTable = doc.add_table(rows=1,cols=3)
menuTable.style= 'TableGrid'
hdr_Cells = menuTable.rows[0].cells
hdr_Cells[0].text = 'Указатель'
hdr_Cells[1].text = 'Значения p'
hdr_Cells[2].text = 'Значения β'


for A in data2.keys():
    row_Cells = menuTable.add_row().cells
    row_Cells[0].text= str(A)
    row_Cells[1].text = str(data2[A][0])
    row_Cells[2].text = str(data2[A][1])



# Сортировка
Dict = collections.OrderedDict(sorted(Dict.items()))

c = list(Dict.keys())
B = list(Dict.values())


Dict2 = collections.OrderedDict(reversed(sorted(Dict2.items())))

c2 = list(Dict2.keys())
B2 = list(Dict2.values())

s = 0
S = []
for i in range(n):
    s += B[i][0]
    S.append(s)

s = 0
S2 = []
for i in range(n):
    s += B2[i][0]
    S2.append(s)


for i in range(n):
    plt.plot(c[i], S[i],
             marker='o',
             markersize=5,
             color='red')

Dict = {}
for i in range(n):
    Dict['A{}'.format(B[i][1])] = ['{:.3f}'.format(c[i]),
                                   '{:.3f}'.format(B[i][0]),
                                   '{:.3f}'.format(S[i])
                                  ]

for i in range(n):
    plt.plot(c2[i], S2[i],
             marker='o',
             markersize=5,
             color='blue')

Dict2 = {}
for i in range(n):
    Dict2['B{}'.format(B2[i][1])] = ['{:.3f}'.format(c2[i]),
                                   '{:.3f}'.format(B2[i][0]),
                                   '{:.3f}'.format(S2[i])
                                  ]

func_val = c[0]
c = [func_val] + c
S = [0] + S
plt.step(c, S, where='post', color='red')


func_val2 = c2[0]
c2 = [func_val2] + c2
S2 = [0] + S2
plt.step(c2, S2, where='post', color='blue')

x = np.arange(0.5, 3, 0.1)

plt.xticks(x,fontsize=8)
plt.xlabel('Цена, p')
plt.ylabel('Объемы, b')
plt.savefig('fig.png')
plt.show()



parag = doc.add_paragraph('\n')
parag = doc.add_paragraph('Данные в отсортированном виде')

menuTable = doc.add_table(rows=1,cols=4)
menuTable.style= 'TableGrid'
hdr_Cells = menuTable.rows[0].cells
hdr_Cells[0].text = 'Указатель'
hdr_Cells[1].text = 'Значения p'
hdr_Cells[2].text = 'Значения α'
hdr_Cells[3].text = 'Значения D(p)'


for A in Dict.keys():
    row_Cells = menuTable.add_row().cells
    row_Cells[0].text= str(A)
    row_Cells[1].text = str(Dict[A][0])
    row_Cells[2].text = str(Dict[A][1])
    row_Cells[3].text = str(Dict[A][2])

parag = doc.add_paragraph('     ')

menuTable = doc.add_table(rows=1,cols=4)
menuTable.style= 'TableGrid'
hdr_Cells = menuTable.rows[0].cells
hdr_Cells[0].text = 'Указатель'
hdr_Cells[1].text = 'Значения x'
hdr_Cells[2].text = 'Значения β'
hdr_Cells[3].text = 'Значения S(p)'

for A in Dict2.keys():
    row_Cells = menuTable.add_row().cells
    row_Cells[0].text= str(A)
    row_Cells[1].text = str(Dict2[A][0])
    row_Cells[2].text = str(Dict2[A][1])
    row_Cells[3].text = str(Dict2[A][2])

parag = doc.add_paragraph('\n')
parag = doc.add_paragraph('Визуализация распределения')

pic = doc.add_picture('fig.png')

interval1 = [list(Dict.items())[0][1][0], list(Dict2.items())[0][1][0]]

parag = doc.add_paragraph('\n')
parag = doc.add_paragraph('Расчет цены аукциона')

menuTable = doc.add_table(rows=1,cols=3)
menuTable.style= 'TableGrid'
hdr_Cells = menuTable.rows[0].cells
hdr_Cells[0].text = 'Отрезок'
hdr_Cells[1].text = 'Знак F((a + b) / 2)'
hdr_Cells[2].text = 'Новый отрезок'

N = None

def recursion(interval, D_val, S_val):

    interval = [float(x) for x in interval]

    elements = []

    for A in Dict.keys():
        if interval[0] < float(Dict[A][0]) < interval[1]:
            elements += [float(Dict[A][0])]
    for A in Dict2.keys():
        if interval[0] < float(Dict2[A][0]) < interval[1]:
            elements += [float(Dict2[A][0])]

    if len(elements) == 1:

        doc.add_paragraph('Цена аукциона: p = ' + str(elements[0]))

        s = min([D_val, S_val])

        array_d = []

        for A in Dict.keys():
            if float(Dict[A][2]) <= s:
                array_d += [float(Dict[A][1])]

        array_s = []

        for A in Dict2.keys():
            if float(Dict2[A][2]) <= s:
                array_s += [float(Dict2[A][1])]

        array_s = list(reversed(array_s))


        if sum(array_d) < sum(array_s):
            array_d += [np.round(s - sum(array_d), 3)]
        else:
            array_s = [np.round(s - sum(array_s), 3)] + array_s

        for i in range(n - len(array_d)):
            array_d += [0.00]
        for i in range(n - len(array_s)):
            array_s = [0.00] + array_s



        doc.add_paragraph('s = min(D(p), S(p)) = min({}, {}) = {}'.format(D_val, S_val, s))


        doc.add_paragraph('x = (' + ', '.join([str(x) for x in array_d]) + ")")
        doc.add_paragraph('y = (' + ', '.join([str(x) for x in array_s]) + ")")
        return elements[0], D_val, S_val

    p = np.mean([float(x) for x in interval])



    for A in Dict.keys():
        if p > float(Dict[A][0]):
            D_val = float(Dict[A][2])
            x_val = float(Dict[A][0])
        if p < float(Dict[A][0]):
            break
    for A in Dict2.keys():
        if p < float(Dict2[A][0]):
            S_val = float(Dict2[A][2])
            y_val =  float(Dict2[A][0])
        if p > float(Dict2[A][0]):
            break


    row_Cells = menuTable.add_row().cells
    row_Cells[0].text = '[' + str(interval[0]) + "; " + str(interval[1]) + "]"


    if D_val < S_val:
        row_Cells[1].text = str("< 0")
        row_Cells[2].text = '[' + str(p) + "; " + str(interval[1]) + "]"
        recursion([p, interval[1]], D_val, S_val)
    else:
        row_Cells[1].text = str("> 0")
        row_Cells[2].text = '[' + str(interval[0]) + "; " + str(p) + "]"
        recursion([interval[0], p], D_val, S_val)


D_val = None
x_val = None
S_val = None
y_val = None
recursion(interval1, D_val, S_val)

doc.save('table.docx')

os.system("start table.docx")
